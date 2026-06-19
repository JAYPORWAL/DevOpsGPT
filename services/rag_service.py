import os
from typing import List, Optional
import pypdf
import docx
from openai import OpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from utils.logger import setup_logger

logger = setup_logger("RAGService")

class GeminiOpenAIEmbeddings(Embeddings):
    """
    Custom LangChain Embeddings implementation that wraps OpenAI Client calls
    directly to bypass LangChain's internal 501 Unimplemented headers when using
    Google's OpenAI compatibility layer. Includes retry and exponential backoff.
    """
    def __init__(self, api_key: str, model: str = "gemini-embedding-2"):
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )
        self.model = model

    def _execute_with_retry(self, func, *args, **kwargs):
        import time
        retries = 0
        backoff_times = [2.0, 4.0, 8.0]
        while retries <= 3:
            try:
                kwargs["timeout"] = 30.0
                return func(*args, **kwargs)
            except Exception as e:
                retries += 1
                if retries <= 3:
                    logger.warning(f"Embedding attempt {retries} failed: {e}. Retrying...")
                    time.sleep(backoff_times[retries - 1])
                else:
                    logger.error(f"Embedding failed after 4 attempts: {e}")
                    raise e

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        response = self._execute_with_retry(
            self.client.embeddings.create,
            input=texts,
            model=self.model
        )
        return [item.embedding for item in response.data]

    def embed_query(self, text: str) -> List[float]:
        response = self._execute_with_retry(
            self.client.embeddings.create,
            input=[text],
            model=self.model
        )
        return response.data[0].embedding

class RAGService:
    """
    Implements a Document RAG (Retrieval-Augmented Generation) pipeline.
    Parses PDF, DOCX, and TXT files, chunks text, creates/updates a local FAISS
    vector index, and retrieves context for query augmentation.
    """

    def __init__(self, api_key: Optional[str] = None, data_dir: str = "data"):
        self.api_key = api_key or os.getenv("GOOGLE_API_KEY")
        self.data_dir = data_dir
        self.faiss_dir = os.path.join(data_dir, "faiss_index")
        self.embeddings = None
        self.vector_store = None

        if self.api_key:
            self._initialize_embeddings()
            self._load_vector_store()

    def set_api_key(self, api_key: str) -> None:
        """Sets a new API key and re-initializes embeddings and vector store."""
        self.api_key = api_key
        self._initialize_embeddings()
        self._load_vector_store()

    def _initialize_embeddings(self) -> None:
        """Initializes our custom Gemini embedding wrapper."""
        try:
            if self.api_key:
                self.embeddings = GeminiOpenAIEmbeddings(api_key=self.api_key)
                logger.info("Embeddings successfully initialized using Google Gemini custom wrapper (gemini-embedding-2).")
            else:
                logger.warning("No API key available to initialize embeddings.")
        except Exception as e:
            logger.error(f"Failed to initialize GeminiOpenAIEmbeddings: {e}")

    def _load_vector_store(self) -> None:
        """Loads local FAISS vector store if it exists."""
        if not self.embeddings:
            logger.warning("Embeddings not initialized. Cannot load FAISS store.")
            return

        # FAISS local storage folder typically contains index.faiss and index.pkl
        faiss_file = os.path.join(self.faiss_dir, "index.faiss")
        if os.path.exists(faiss_file):
            try:
                self.vector_store = FAISS.load_local(
                    self.faiss_dir,
                    self.embeddings,
                    allow_dangerous_deserialization=True
                )
                logger.info("FAISS vector store loaded successfully from local storage.")
            except Exception as e:
                logger.error(f"Error loading local FAISS store: {e}", exc_info=True)
        else:
            logger.info("No local FAISS index found. Ready to ingest new documents.")

    def _extract_text(self, filepath: str) -> str:
        """Helper to extract raw text from PDF, DOCX, or TXT files."""
        ext = os.path.splitext(filepath)[1].lower()
        text = ""

        if ext == ".pdf":
            try:
                with open(filepath, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    for idx, page in enumerate(reader.pages):
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                logger.info(f"Extracted PDF text from {filepath}")
            except Exception as e:
                logger.error(f"Failed to parse PDF {filepath}: {e}")
                raise e

        elif ext == ".docx":
            try:
                doc = docx.Document(filepath)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                logger.info(f"Extracted DOCX text from {filepath}")
            except Exception as e:
                logger.error(f"Failed to parse DOCX {filepath}: {e}")
                raise e

        elif ext == ".txt":
            try:
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                logger.info(f"Extracted TXT text from {filepath}")
            except Exception as e:
                logger.error(f"Failed to parse TXT {filepath}: {e}")
                raise e
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        return text

    def ingest_document(self, filepath: str, original_filename: str) -> int:
        """
        Parses a document, chunks the text, computes embeddings, and indexes
        it in the FAISS vector database.
        
        Args:
            filepath: Path to the file on disk.
            original_filename: The human-readable filename for sourcing.
            
        Returns:
            int: Number of chunks indexed.
        """
        if not self.embeddings:
            raise ValueError("Embeddings are not initialized. Please supply a valid GOOGLE_API_KEY.")

        # 1. Extract raw text
        raw_text = self._extract_text(filepath)
        if not raw_text.strip():
            logger.warning(f"No text extracted from file {original_filename}")
            return 0

        # 2. Chunk text
        # standard chunk size 1000 with 200 character overlap
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(raw_text)
        
        # 3. Create LangChain Document objects
        documents = [
            Document(
                page_content=chunk,
                metadata={"source": original_filename, "chunk_index": i}
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # 4. Add to Vector Store
        if self.vector_store is None:
            self.vector_store = FAISS.from_documents(documents, self.embeddings)
            logger.info("Created new FAISS index.")
        else:
            self.vector_store.add_documents(documents)
            logger.info("Added documents to existing FAISS index.")

        # 5. Save locally
        try:
            self.vector_store.save_local(self.faiss_dir)
            logger.info("FAISS vector store saved locally.")
        except Exception as e:
            logger.error(f"Failed to save FAISS store: {e}")
            raise e

        return len(documents)

    def similarity_search(self, query: str, k: int = 3) -> List[Document]:
        """
        Queries the vector index for similar content blocks.
        
        Args:
            query: User's query string.
            k: Top k chunks to return.
            
        Returns:
            List[Document]: List of LangChain Document chunks.
        """
        if self.vector_store is None:
            logger.warning("Vector store is not initialized. No files have been indexed yet.")
            return []
            
        try:
            results = self.vector_store.similarity_search(query, k=k)
            logger.info(f"Similarity search retrieved {len(results)} chunks.")
            return results
        except Exception as e:
            logger.error(f"Similarity search failed: {e}")
            return []

    def clear_vector_store(self) -> None:
        """Removes local index files and resets store."""
        self.vector_store = None
        for file in ["index.faiss", "index.pkl"]:
            filepath = os.path.join(self.faiss_dir, file)
            if os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except Exception as e:
                    logger.warning(f"Could not remove {file}: {e}")
        logger.info("Vector store index cleared.")
